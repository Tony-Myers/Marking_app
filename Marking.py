# Standard library imports
import streamlit as st
from openai import OpenAI
import pandas as pd
import docx
from PyPDF2 import PdfReader
from io import BytesIO, StringIO
import tiktoken
import csv
import re
import os

# Set your OpenAI API key and password from secrets
PASSWORD = st.secrets["password"]
OPENAI_API_KEY = st.secrets["openai_api_key"]

# Instantiate the OpenAI client
client = OpenAI(api_key=OPENAI_API_KEY)

# Initialize the tiktoken encoder for GPT-4
try:
    encoding = tiktoken.encoding_for_model("gpt-4")
except KeyError:
    encoding = tiktoken.get_encoding("cl100k_base")

MAX_TOKENS = 7000  # Maximum tokens for GPT-4
PROMPT_BUFFER = 1000  # Buffer to ensure we don't exceed the limit

def count_tokens(text, encoding):
    """Counts the number of tokens in a given text."""
    return len(encoding.encode(text))

def truncate_text(text, max_tokens, encoding):
    """Truncates text to fit within max_tokens."""
    tokens = encoding.encode(text)
    if len(tokens) > max_tokens:
        truncated_tokens = tokens[:max_tokens]
        return encoding.decode(truncated_tokens)
    return text

def call_chatgpt(prompt, model="gpt-4", max_tokens=3000, temperature=0.3, retries=2):
    """Calls the OpenAI API using the client instance and returns the response as text."""
    for attempt in range(retries):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens,
                temperature=temperature,
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            if attempt < retries - 1:
                continue
            else:
                st.error(f"API Error: {e}")
                return None

def check_password():
    """Prompts the user for a password and checks it."""
    def password_entered():
        if st.session_state["password"] == PASSWORD:
            st.session_state["password_correct"] = True
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.text_input("Enter the password", type="password", on_change=password_entered, key="password")
        return False
    elif not st.session_state["password_correct"]:
        st.text_input("Enter the password", type="password", on_change=password_entered, key="password")
        st.error("😕 Password incorrect")
        return False
    else:
        return True

def parse_csv_section(csv_text):
    """Parses a CSV section line by line, handling quoted fields."""
    try:
        # Use StringIO to read the CSV text
        csv_io = StringIO(csv_text)
        # Read the CSV with proper quoting
        df = pd.read_csv(csv_io, dtype={'Criterion': str, 'Score': float}, quotechar='"', skipinitialspace=True)
        return df
    except Exception as e:
        st.error(f"Error parsing CSV: {e}")
        return None

def summarize_text(text):
    """Summarizes the given text using OpenAI API."""
    summary_prompt = f"""
You are an assistant that summarizes academic papers. Please provide a concise summary (max 500 words) of the following text:

{text}
"""
    summary = call_chatgpt(summary_prompt, max_tokens=800, temperature=0.3)
    return summary if summary else text  # Fallback to original text if summarization fails

def extract_weight(criterion_name):
    """
    Extracts the weight from the criterion name.
    For example, from "Linking Theory to Issue (15%)", it extracts 15.0
    """
    match = re.search(r'\((\d+)%\)', criterion_name)
    if match:
        return float(match.group(1))
    else:
        return 0.0  # Default weight if not found

def initialize_session_state():
    """Initializes session state for storing feedback."""
    if 'feedbacks' not in st.session_state:
        st.session_state['feedbacks'] = {}

# Function to extract text from .pdf files
def extract_text_from_pdf(pdf_file):
    pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
    text = ""
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        text += page.get_text()
    return text

# Function to extract text from .txt files
def extract_text_from_txt(txt_file):
    return txt_file.read().decode("utf-8")

# Streamlit file uploader
uploaded_files = st.file_uploader("Upload files", type=['pdf','docx','txt'], accept_multiple_files=True)

if uploaded_files:
    for uploaded_file in uploaded_files:
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        try:
            if file_extension == '.docx':
                text = extract_text_from_docx(uploaded_file)
            elif file_extension == '.pdf':
                text = extract_text_from_pdf(uploaded_file)
            elif file_extension == '.txt':
                text = extract_text_from_txt(uploaded_file)
            else:
                st.error(f"Unsupported file type: {uploaded_file.name}")
                continue
            st.write(f"Extracted text from {uploaded_file.name}:")
            st.write(text)
        except Exception as e:
            st.error(f"Error reading {uploaded_file.name}: {e}")

# Function to extract text from .docx files
def extract_text_from_docx(docx_file):
    doc = docx.Document(docx_file)
    return '\n'.join([para.text for para in doc.paragraphs])

# Function to extract text from .pdf files
def extract_text_from_pdf(pdf_file):
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

# Function to extract text from .txt files
def extract_text_from_txt(txt_file):
    return txt_file.read().decode("utf-8")
    st.error(f"Error reading {uploaded_file.name}: {e}")

if uploaded_files:
    for uploaded_file in uploaded_files:
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        if file_extension == '.docx':
            try:
                text = extract_text_from_docx(uploaded_file)
                st.write(f"Extracted text from {uploaded_file.name}:")
                st.write(text)
            except Exception as e:
                st.error(f"Error reading {uploaded_file.name}: {e}")
        elif file_extension == '.pdf':
            try:
                text = extract_text_from_pdf(uploaded_file)
                st.write(f"Extracted text from {uploaded_file.name}:")
                st.write(text)
            except Exception as e:
                st.error(f"Error reading {uploaded_file.name}: {e}")
        elif file_extension == '.txt':
            try:
                text = extract_text_from_txt(uploaded_file)
                st.write(f"Extracted text from {uploaded_file.name}:")
                st.write(text)
            except Exception as e:
                st.error(f"Error reading {uploaded_file.name}: {e}")
        else:
            st.error(f"Unsupported file type: {uploaded_file.name}")

def main():
    initialize_session_state()
    
    if check_password():
        st.title("✏️ Automated Assignment Grading and Feedback © Tony Myers")

        # Define the column name for criterion here to make it accessible throughout the function
        criterion_column = 'Criterion'

        st.header("Assignment Task and Level of Work")
        assignment_task = st.text_area("Enter the Assignment Task or Instructions, as well as the HE level the work should be marked at", height=150)

        st.header("Upload Files")
        rubric_file = st.file_uploader("Upload Grading Rubric (CSV)", type=['csv'])
        submissions = st.file_uploader("Upload Student Submissions (.docx, .pdf, .txt)", type=['docx', 'pdf', 'txt'], accept_multiple_files=True)

        if rubric_file and submissions:
            if st.button("Run Marking"):
                # Read the grading rubric
                try:
                    original_rubric_df = pd.read_csv(rubric_file, dtype={criterion_column: str})
                except Exception as e:
                    st.error(f"Error reading rubric: {e}")
                    return

                if criterion_column not in original_rubric_df.columns:
                    st.error(f"Rubric must contain a '{criterion_column}' column.")
                    return

                # Ensure Criterion column is string type for consistency in both dataframes
                original_rubric_df[criterion_column] = original_rubric_df[criterion_column].astype(str)

                # Extract Weight from Criterion and clean Criterion names
                original_rubric_df['Weight'] = original_rubric_df[criterion_column].apply(extract_weight)
                original_rubric_df[criterion_column] = original_rubric_df[criterion_column].apply(lambda x: re.sub(r'\s*\(\d+%\)', '', x))

                # Generate rubric CSV string with all fields quoted
                rubric_csv_string = original_rubric_df.to_csv(index=False, quoting=csv.QUOTE_ALL)

                # Get the list of percentage range columns (e.g., '0-59%', '60
                # Get the list of percentage range columns (e.g., '0-59%', '60-69%', etc.)
                percentage_columns = [col for col in original_rubric_df.columns if '%' in col]

                # Get the list of criteria
                criteria_list = original_rubric_df[criterion_column].tolist()
                criteria_string = '\n'.join(criteria_list)

                for submission in submissions:
                    student_name = os.path.splitext(submission.name)[0]
                    
                    # Skip if feedback already exists for this student
                    if student_name in st.session_state['feedbacks']:
                        st.info(f"Feedback already generated for {student_name}.")
                        continue

                    st.header(f"Processing {student_name}'s Submission")

                    # Read student submission
                    try:
                        doc = docx.Document(submission)
                        student_text = '\n'.join([para.text for para in doc.paragraphs])
                    except Exception as e:
                        st.error(f"Error reading submission {submission.name}: {e}")
                        continue

                    # Summarize the student submission if it's too long
                    student_tokens = count_tokens(student_text, encoding)
                    max_submission_tokens = MAX_TOKENS - PROMPT_BUFFER  # Reserve tokens for prompt and other texts

                    if student_tokens > (MAX_TOKENS * 0.6):  # If submission is more than ~4k tokens
                        st.info(f"Summarizing {student_name}'s submission to reduce token count.")
                        student_text = summarize_text(student_text)
                        if not student_text:
                            st.error(f"Failed to summarize submission for {student_name}.")
                            continue

                    # Recalculate tokens after summarization
                    student_tokens = count_tokens(student_text, encoding)

                    # Prepare prompt for ChatGPT with modifications
                    prompt = f"""
You are an experienced UK academic tasked with grading a student's assignment based on the provided rubric and assignment instructions. Please ensure that your feedback adheres to UK Higher Education standards for undergraduate work, noting the level provided by the user. Use British English spelling throughout your feedback.

**Instructions:**

- Review the student's submission thoroughly and be strict in applying the criteria.
- For **each criterion** in the list below, assign a numerical score between 0 and 100 (e.g., 75) and provide a brief but nuanced comment.
- Ensure that the score is numeric without any extra symbols or text.
- The scores should reflect the student's performance according to the descriptors in the rubric.
- **Be strict in your grading to align with UK undergraduate standards.**
- **Assess the quality of writing and referencing style, ensuring adherence to the 'Cite them Right' guidelines (2008, Pear Tree Books). Provide a brief comment on these aspects in the overall comments but refer to the referencing style as Birmingham Newman University’s referencing style in feedback.**

**List of Criteria:**
{criteria_string}

**Rubric (in CSV format):**
{rubric_csv_string}

**Assignment Task:**
{assignment_task}

**Student's Submission:**
{student_text}

**Your Output Format:**

Please output your feedback in the exact format below, ensuring you include **all criteria**:

**Important Notes:**

- **Begin your response with the CSV section**, starting with "Criterion,Score,Comment".
- **Include all criteria** from the list provided.
- **Do not omit any sections**.
- **Do not include any additional text** or formatting outside of the specified format.
- **Ensure that 'Overall Comments:' and 'Feedforward:' are included exactly as shown**, with the colon and on separate lines.
- **Do not use markdown formatting** like bold, italics, or headers.
- **Ensure there are no extra lines or missing lines**.
- **Your entire response should be in plain text**.
- **Use second person narrative ("You...") in Overall Comments and Feedforward.**
- **Ensure British English spelling is used.**
- **Overall Comments should not exceed 150 words.**
- **Feedforward should be a bulleted list within 150 words.**
- **Comments on each criterion should be concise, in the second person, and not exceed 150 words in total.**
- **Ensure that all fields containing commas are enclosed in double quotes.**
- **Provide an example of the expected CSV format below.**

**Example:**

Overall Comments:
Your essay provides a solid foundation in linking Functionalism and Critical Theory to the issue of racism in football. While you demonstrate a good understanding of the theories, the analysis could be more critical and less descriptive. The structure of your essay is generally clear, but some sections could benefit from improved coherence and flow. Your referencing style is mostly consistent with the 'Cite them Right' guidelines, but ensure all sources are correctly formatted and cited.

Feedforward:
- Focus on making explicit connections between theory and issue to strengthen your analysis.
- Aim for a more critical application of theory, moving beyond description to provide deeper insights.
- Enhance the structure by ensuring each paragraph has a clear topic sentence and logical progression.
- Review 'Cite them Right' guidelines to ensure all references are correctly formatted.
- Consider using more varied sources to support your arguments and provide a broader perspective.

**Note:** Additionally, please include a **Total Mark** based on the weighted scores of each criterion. This total mark should only appear in the downloaded `.docx` file and not in the Streamlit app.

                    """

                    # Estimate total tokens
                    total_tokens = count_tokens(prompt, encoding)
                    if total_tokens > MAX_TOKENS:
                        st.error(f"The prompt for {student_name} exceeds the maximum token limit. Please reduce the length of the rubric, assignment instructions, or the student submission.")
                        continue

                    # Calculate max tokens for response
                    max_response_tokens = MAX_TOKENS - total_tokens

                    # Call ChatGPT API
                    feedback = call_chatgpt(prompt, max_tokens=max_response_tokens, temperature=0.3)
                    if feedback:
                        st.success(f"Feedback generated for {student_name}")

                        # Display AI's response for debugging (optional)
                        # st.text_area("AI Response", feedback, height=300)

                        # Parse the feedback
                        try:
                            # Check if 'Overall Comments:' is in the feedback
                            if 'Overall Comments:' in feedback:
                                # Split the feedback into CSV and comments sections
                                csv_feedback = feedback.split('Overall Comments:', 1)[0].strip()
                                comments_section = feedback.split('Overall Comments:', 1)[1].strip()
                            else:
                                st.error("The AI's response is missing 'Overall Comments:'. Please adjust the prompt or try again.")
                                st.write("AI Response:")
                                st.code(feedback)
                                continue

                            # Check if the CSV section is present
                            if 'Criterion,Score,Comment' in csv_feedback:
                                completed_rubric_df = parse_csv_section(csv_feedback)
                                if completed_rubric_df is None:
                                    st.error("Failed to parse the CSV feedback.")
                                    st.write("AI Response:")
                                    st.code(feedback)
                                    continue
                            else:
                                st.error("The AI's response is missing the CSV section with 'Criterion,Score,Comment'. Please adjust the prompt or try again.")
                                st.write("AI Response:")
                                st.code(feedback)
                                continue

                            # Ensure all criteria are present
                                                       # Ensure all criteria are present
                            missing_criteria = set(original_rubric_df[criterion_column]) - set(completed_rubric_df[criterion_column])
                            if missing_criteria:
                                st.warning(f"The AI feedback is missing the following criteria: {missing_criteria}")

                            # Extract overall comments and feedforward
                            overall_comments = ''
                            feedforward = ''
                            if 'Feedforward:' in comments_section:
                                overall_comments = comments_section.split('Feedforward:', 1)[0].strip()
                                feedforward = comments_section.split('Feedforward:', 1)[1].strip()
                            else:
                                st.warning("The AI's response is missing 'Feedforward:'. Please adjust the prompt or try again.")
                                overall_comments = comments_section.strip()
                                feedforward = ''

                            # Merge the dataframes with specified suffixes
                            merged_rubric_df = original_rubric_df.merge(
                                completed_rubric_df[[criterion_column, 'Score', 'Comment']],
                                on=criterion_column,
                                how='left',
                                suffixes=('', '_ai')  # Suffix for AI feedback columns
                            ).dropna(how="all", axis=1)

                            # Calculate Weighted Scores and Total Mark
                            merged_rubric_df['Weighted_Score'] = merged_rubric_df['Weight'] * merged_rubric_df['Score'] / 100
                            total_mark = merged_rubric_df['Weighted_Score'].sum()

                            # Store the feedback in session state to persist across reruns
                            st.session_state['feedbacks'][student_name] = {
                                'merged_rubric_df': merged_rubric_df,
                                'overall_comments': overall_comments,
                                'feedforward': feedforward,
                                'percentage_columns': percentage_columns,
                                'total_mark': total_mark  # Store total_mark here
                            }

                            # Optionally, display DataFrames for debugging
                            # st.write("Completed Rubric DataFrame:")
                            # st.dataframe(completed_rubric_df)

                            # st.write("Merged Rubric DataFrame:")
                            # st.dataframe(merged_rubric_df)

                        except Exception as e:
                            st.error(f"Error parsing AI response: {e}")
                            st.write("AI Response:")
                            st.code(feedback)
                            continue

            st.success("All submissions have been processed.")

        # After processing, provide download buttons without displaying any feedback
            if st.session_state.get('feedbacks'):
                st.header("Generated Feedbacks")
                for student_name, feedback_data in st.session_state['feedbacks'].items():
                    # No display of rubric scores or comments in the app

                    # Create Word document for feedback
                    feedback_doc = docx.Document()

                    # Set page to landscape
                    section = feedback_doc.sections[0]
                    section.orientation = WD_ORIENT.LANDSCAPE
                    new_width, new_height = section.page_height, section.page_width
                    section.page_width = new_width
                    section.page_height = new_height

                    feedback_doc.add_heading(f"Feedback for {student_name}", level=1)

                    if not feedback_data['merged_rubric_df'].empty:
                        # Prepare columns for the Word table
                        table_columns = [criterion_column] + feedback_data['percentage_columns'] + ['Score', 'Comment']
                        table = feedback_doc.add_table(rows=1, cols=len(table_columns))
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        for i, column in enumerate(table_columns):
                            hdr_cells[i].text = str(column)

                        # Add data rows and apply shading to the appropriate descriptor cell
                        for _, row in feedback_data['merged_rubric_df'].iterrows():
                            row_cells = table.add_row().cells
                            score = row['Score']
                            for i, col_name in enumerate(table_columns):
                                cell = row_cells[i]
                                cell_text = str(row[col_name])
                                cell.text = cell_text

                                # Apply shading to the descriptor cell matching the score range
                                if col_name in feedback_data['percentage_columns'] and pd.notnull(score):
                                    # Extract numeric values from the percentage range
                                    range_text = col_name.replace('%', '').strip()
                                    lower_upper = range_text.split('-')
                                    if len(lower_upper) == 2:
                                        try:
                                            lower = float(lower_upper[0].strip())
                                            upper = float(lower_upper[1].strip())

                                            # Use the score as a float
                                            score_value = float(score)

                                            if lower <= score_value <= upper:
                                                # Apply green shading to this cell
                                                shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
                                                cell._tc.get_or_add_tcPr().append(shading_elm)
                                        except ValueError as e:
                                            st.warning(f"Error converting score or range to float: {e}")
                                            continue

                    # Add overall comments and feedforward
                    feedback_doc.add_heading('Overall Comments', level=2)
                    feedback_doc.add_paragraph(feedback_data['overall_comments'].strip())
                    feedback_doc.add_heading('Feedforward', level=2)
                    feedback_doc.add_paragraph(feedback_data['feedforward'].strip())

                    # Add Total Mark
                    feedback_doc.add_heading('Total Mark', level=2)
                    feedback_doc.add_paragraph(f"{feedback_data['total_mark']:.2f}")

                    buffer = BytesIO()
                    feedback_doc.save(buffer)
                    buffer.seek(0)

                    # Provide download button for the feedback document
                    st.download_button(
                        label=f"Download Feedback for {student_name}",
                        data=buffer,
                        file_name=f"{student_name}_feedback.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

if __name__ == "__main__":
    main()


 

 
